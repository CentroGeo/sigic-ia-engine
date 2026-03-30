import io
import os
import re

from docx import Document
from docx.shared import Pt


def _get_report_font_name() -> str:
    """Extrae el nombre de fuente primario de REPORT_FONT (sin comillas ni fallbacks CSS)."""
    raw = os.environ.get("REPORT_FONT", '"Montserrat", sans-serif')
    primary = raw.split(",")[0].strip().strip('"').strip("'")
    return primary


def render_docx(content: str, output_format: str, use_letterhead: bool = False) -> bytes:
    """
    Convierte *content* a bytes DOCX.

    Parameters
    ----------
    content       : texto generado por el LLM
    output_format : 'markdown' | 'plain_text'
    use_letterhead: booleano para aplicar formato SECIHTI
    """
    # Limpiar siempre: quitar code fences que el LLM a veces agrega
    content = _strip_code_fences(content)

    doc = None
    if use_letterhead:
        import os
        from django.conf import settings
        template_path = os.path.join(settings.BASE_DIR, "reports", "templates", "plantilla_secihti.docx")
        if os.path.exists(template_path):
            doc = Document(template_path)
        else:
            print(f"[REPORT] ADVERTENCIA: Plantilla {template_path} no encontrada.")
            doc = Document()
    else:
        doc = Document()

    font_name = _get_report_font_name()

    # Siempre parsear como markdown: el LLM suele generar markdown incluso
    # cuando se solicita plain_text. El parser lo maneja en ambos casos.
    _parse_markdown_to_docx(doc, content, font_name)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Helpers internos
# ---------------------------------------------------------------------------

def _apply_font(paragraph, font_name: str) -> None:
    """Aplica font_name a todos los runs de un párrafo."""
    if not font_name:
        return
    for run in paragraph.runs:
        run.font.name = font_name


def _add_paragraph_with_font(doc: Document, text: str, font_name: str):
    p = doc.add_paragraph(text)
    _apply_font(p, font_name)
    return p


def _parse_markdown_to_docx(doc: Document, content: str, font_name: str = "") -> None:
    """Convierte Markdown básico a párrafos/headings de docx."""
    lines = content.splitlines()
    in_table = False
    table_rows: list[list[str]] = []

    for line in lines:
        # Heading 1
        if line.startswith("# "):
            _flush_table(doc, table_rows, font_name)
            table_rows = []
            in_table = False
            p = doc.add_heading(line[2:].strip(), level=1)
            _apply_font(p, font_name)

        # Heading 2
        elif line.startswith("## "):
            _flush_table(doc, table_rows, font_name)
            table_rows = []
            in_table = False
            p = doc.add_heading(line[3:].strip(), level=2)
            _apply_font(p, font_name)

        # Heading 3
        elif line.startswith("### "):
            _flush_table(doc, table_rows, font_name)
            table_rows = []
            in_table = False
            p = doc.add_heading(line[4:].strip(), level=3)
            _apply_font(p, font_name)

        # Heading 4+ → nivel 3 (docx no tiene h4 por defecto)
        elif re.match(r"^#{4,} ", line):
            _flush_table(doc, table_rows, font_name)
            table_rows = []
            in_table = False
            text = re.sub(r"^#+\s+", "", line).strip()
            p = doc.add_heading(text, level=3)
            _apply_font(p, font_name)

        # Table row
        elif line.startswith("|"):
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            # Ignore separator rows like |---|---|
            if all(re.match(r"^[-: ]+$", c) for c in cells if c):
                continue
            table_rows.append(cells)
            in_table = True

        # Bullet list
        elif line.startswith("- ") or line.startswith("* "):
            _flush_table(doc, table_rows, font_name)
            table_rows = []
            in_table = False
            clean = _strip_inline_md(line[2:].strip())
            _add_paragraph_with_font(doc, f"• {clean}", font_name)

        # Numbered list
        elif re.match(r"^\d+\. ", line):
            _flush_table(doc, table_rows, font_name)
            table_rows = []
            in_table = False
            clean = _strip_inline_md(line.strip())
            _add_paragraph_with_font(doc, clean, font_name)

        # Empty line — flush pending table if any
        elif line.strip() == "":
            if in_table:
                _flush_table(doc, table_rows, font_name)
                table_rows = []
                in_table = False
            # else just skip blank lines

        # Regular paragraph — strip inline markdown bold/italic
        else:
            _flush_table(doc, table_rows, font_name)
            table_rows = []
            in_table = False
            clean = _strip_inline_md(line)
            if clean:
                _add_paragraph_with_font(doc, clean, font_name)

    # Flush any remaining table
    _flush_table(doc, table_rows, font_name)


def _parse_plain_text_to_docx(doc: Document, content: str) -> None:
    """Divide por doble salto de línea y agrega un párrafo por bloque."""
    blocks = content.split("\n\n")
    for block in blocks:
        block = block.strip()
        if not block:
            continue
        # Lines in ALL CAPS treated as headings
        first_line = block.splitlines()[0].strip()
        if first_line == first_line.upper() and len(first_line) > 3:
            doc.add_heading(first_line, level=2)
            rest = "\n".join(block.splitlines()[1:]).strip()
            if rest:
                doc.add_paragraph(rest)
        else:
            doc.add_paragraph(block)


def _flush_table(doc: Document, rows: list[list[str]], font_name: str = "") -> None:
    if not rows:
        return
    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=max_cols)
    table.style = "Table Grid"
    for i, row_data in enumerate(rows):
        row = table.add_row()
        for j, cell_text in enumerate(row_data):
            if j < max_cols:
                row.cells[j].text = cell_text
                for run in row.cells[j].paragraphs[0].runs:
                    if i == 0:
                        run.bold = True
                    if font_name:
                        run.font.name = font_name


def _strip_code_fences(content: str) -> str:
    """Quita bloques ```...``` o ```markdown...``` que el LLM envuelve a veces."""
    content = content.strip()
    content = re.sub(r"^```[a-zA-Z]*\n?", "", content)
    content = re.sub(r"\n?```$", "", content)
    return content.strip()


def _strip_inline_md(text: str) -> str:
    """Remove bold/italic markers from a line."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"__(.+?)__", r"\1", text)
    text = re.sub(r"_(.+?)_", r"\1", text)
    return text
