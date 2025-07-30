from rest_framework.response import Response
from django.http import JsonResponse
from rest_framework.decorators import api_view
from django.views.decorators.csrf import csrf_exempt
from django.db.models import Q
from fileuploads.models import Workspace, Context, Files
from django.db import transaction
import pandas as pd
from docx import Document as DocxDocument
from pdfminer.high_level import extract_text as extract_pdf_text
from rest_framework.views import APIView
from rest_framework import status
from llama_index.llms.ollama import Ollama
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.vector_stores.postgres import PGVectorStore
from llama_index.core import VectorStoreIndex, StorageContext
from llama_index.core.schema import Document
import  json 
import os
import shutil

"""
    Secciones de apis para workspaces
"""
@api_view(["GET", "POST"])
def list_workspaces(request):
    user_id = request.GET.get("user_id")
    
    list_workspaces = list(Workspace.objects.filter(
        Q(user_id=user_id) | Q(public=True),
        active=True
    ).values(
        'id', 'title', 'description', 'user_id', 'active', 'public', 'created_date', 'image_type'
    ))
    
    
    return JsonResponse(list(list_workspaces), safe=False)


@api_view(["GET", "POST"])
def list_admin_workspaces(request):
    user_id = request.GET.get("user_id")
    
    list_workspaces = list(Workspace.objects.filter(
        user_id=user_id,
    ).values(
        'id', 'title', 'description', 'user_id', 'active', 'public', 'created_date', 'image_type'
    ))
    
    
    return JsonResponse(list(list_workspaces), safe=False)

@api_view(["GET", "POST"])
@csrf_exempt
def create_admin_workspaces(request):
    user_id = request.GET.get("user_id")
    workspace_data = request.POST.copy()
    answer = {
        "id": None,
        "saved": False
    }
    
    file = request.FILES.get('file', None)
    if file:
        filename = file.name
        file_extension = filename.split('.')[-1].lower()
        
        if file_extension not in ['jpg', 'jpeg', 'png']:
            return JsonResponse(
                {"error": "El archivo debe ser una imagen (jpg, jpeg, png)"},
                status=400
            )
    
    try:
        with transaction.atomic():
            
            new_workspace               = Workspace()
            new_workspace.user_id       = user_id
            new_workspace.title         = workspace_data.get("title")
            new_workspace.description   = workspace_data.get("description")
            new_workspace.public        = workspace_data.get("public")
            
            if file:   
                new_workspace.image_type = file_extension
                
            new_workspace.save()
            
            if file:   
                try:
                    # Crear el directorio si no existe
                    upload_dir = "uploaded_images"
                    if not os.path.exists(upload_dir):
                        os.makedirs(upload_dir)

                    workspaces_dir = os.path.join(upload_dir, "workspaces")
                    if not os.path.exists(workspaces_dir):
                        os.makedirs(workspaces_dir)
                    
                    # Guardar la imagen en el directorio
                    file_path = os.path.join(workspaces_dir, '{}.{}'.format(new_workspace.id, file_extension))
                    with open(file_path, "wb") as f:
                        shutil.copyfileobj(file.file, f)
                except:
                    l = 0
                    
            answer["id"] = new_workspace.id
            answer["saved"] = True
        return JsonResponse(answer, status=200)
    except Exception as e:
        return JsonResponse({"status": "error", "message": str(e)}, status=400)


@api_view(["GET", "POST"])
@csrf_exempt
def edit_admin_workspaces(request, workspace_id):
    user_id = request.GET.get("user_id")
    workspace_data = request.POST.copy()
    answer = {
        "id": None,
        "saved": False
    }
    
    file = request.FILES.get('file', None)
    if file:
        filename = file.name
        file_extension = filename.split('.')[-1].lower()
        
        if file_extension not in ['jpg', 'jpeg', 'png']:
            return JsonResponse(
                {"error": "El archivo debe ser una imagen (jpg, jpeg, png)"},
                status=400
            )
    
    try:
        with transaction.atomic():
            
            get_workspace               = Workspace.objects.get(id=workspace_id)
            get_workspace.user_id       = user_id
            get_workspace.title         = workspace_data.get("title")
            get_workspace.description   = workspace_data.get("description")
            get_workspace.public        = workspace_data.get("public")
            
            if file:   
                get_workspace.image_type = file_extension
                
            get_workspace.save()
            
            if file:   
                try:
                    # Crear el directorio si no existe
                    upload_dir = "uploaded_images"
                    if not os.path.exists(upload_dir):
                        os.makedirs(upload_dir)

                    workspaces_dir = os.path.join(upload_dir, "workspaces")
                    if not os.path.exists(workspaces_dir):
                        os.makedirs(workspaces_dir)
                    
                    # Guardar la imagen en el directorio
                    file_path = os.path.join(workspaces_dir, '{}.{}'.format(workspace_id, file_extension))
                    with open(file_path, "wb") as f:
                        shutil.copyfileobj(file.file, f)
                except:
                    l = 0
                    
            answer["id"] = get_workspace.id
            answer["saved"] = True
        return JsonResponse(answer, status=200)
    except Exception as e:
        return JsonResponse({"status": "error", "message": str(e)}, status=400)



"""
    Secciones de apis para contexts
"""
@api_view(["GET", "POST"])
def list_workspaces_contexts(request, workspace_id):
    user_id = request.GET.get("user_id")
    
    list_workspaces = list(Context.objects.filter(
        Q(user_id=user_id) | Q(public=True),
        active=True,
        workspace_id=workspace_id
    ).values(
        'id', 'title', 'description', 'user_id', 'active', 'public', 'created_date', 'image_type'
    ))
    
    
    return JsonResponse(list(list_workspaces), safe=False)


@api_view(["GET", "POST"])
def list_admin_workspaces_contexts(request, workspace_id):
    user_id = request.GET.get("user_id")
    
    list_workspaces_contexts = list(Context.objects.filter(
        user_id=user_id,
        workspace_id=workspace_id
    ).values(
        'id', 'title', 'description', 'user_id', 'active', 'public', 'created_date', 'image_type'
    ))
    
    
    return JsonResponse(list(list_workspaces_contexts), safe=False)

@api_view(["GET", "POST"])
def create_admin_workspaces_contexts(request):
    user_id = request.GET.get("user_id")
    context_data = request.POST.copy()
    
    answer = {
        "id": None,
        "saved": False
    }
    
    file = request.FILES.get('file', None)
    if file:
        filename = file.name
        file_extension = filename.split('.')[-1].lower()
        
        if file_extension not in ['jpg', 'jpeg', 'png']:
            return JsonResponse(
                {"error": "El archivo debe ser una imagen (jpg, jpeg, png)"},
                status=400
            )
    
    try:
        with transaction.atomic():
            getWorkspace = Workspace.objects.get(id=context_data.get("workspace_id"))
            
            new_context               = Context()
            new_context.workspace     = getWorkspace
            new_context.user_id       = user_id
            new_context.title         = context_data.get("title")
            new_context.description   = context_data.get("description")
            new_context.public        = context_data.get("public")
            
            if file:   
                new_context.image_type = file_extension
                
            new_context.save()
            
            if file:   
                try:
                    # Crear el directorio si no existe
                    upload_dir = "uploaded_images"
                    if not os.path.exists(upload_dir):
                        os.makedirs(upload_dir)

                    contexts_dir = os.path.join(upload_dir, "contexts")
                    if not os.path.exists(contexts_dir):
                        os.makedirs(contexts_dir)
                    
                    # Guardar la imagen en el directorio
                    file_path = os.path.join(contexts_dir, '{}.{}'.format(new_context.id, file_extension))
                    with open(file_path, "wb") as f:
                        shutil.copyfileobj(file.file, f)
                except:
                    l = 0
                    
            answer["id"] = new_context.id
            answer["saved"] = True
        
        return JsonResponse(answer, status=200)    
    except Exception as e:
        return JsonResponse({"status": "error", "message": str(e)}, status=400)

@api_view(["GET", "POST"])
def edit_admin_workspaces_contexts(request, context_id):
    user_id = request.GET.get("user_id")
    context_data = request.POST.copy()
    
    answer = {
        "id": None,
        "saved": False
    }
    
    file = request.FILES.get('file', None)
    if file:
        filename = file.name
        file_extension = filename.split('.')[-1].lower()
        
        if file_extension not in ['jpg', 'jpeg', 'png']:
            return JsonResponse(
                {"error": "El archivo debe ser una imagen (jpg, jpeg, png)"},
                status=400
            )
    
    try:
        with transaction.atomic():
            get_context               = Context.objects.get(id=context_id)
            get_context.user_id       = user_id
            get_context.title         = context_data.get("title")
            get_context.description   = context_data.get("description")
            get_context.public        = context_data.get("public")
            
            if file:   
                get_context.image_type = file_extension
                
            get_context.save()
            
            if file:   
                try:
                    # Crear el directorio si no existe
                    upload_dir = "uploaded_images"
                    if not os.path.exists(upload_dir):
                        os.makedirs(upload_dir)

                    contexts_dir = os.path.join(upload_dir, "contexts")
                    if not os.path.exists(contexts_dir):
                        os.makedirs(contexts_dir)
                    
                    # Guardar la imagen en el directorio
                    file_path = os.path.join(contexts_dir, '{}.{}'.format(context_id, file_extension))
                    with open(file_path, "wb") as f:
                        shutil.copyfileobj(file.file, f)
                except:
                    l = 0
                    
            answer["id"] = get_context.id
            answer["saved"] = True
        
        return JsonResponse(answer, status=200)    
    except Exception as e:
        return JsonResponse({"status": "error", "message": str(e)}, status=400)

"""
    Secciones de apis para files
"""
@api_view(["GET", "POST"])
def list_admin_workspaces_contexts_files(request, workspace_id, context_id):
    user_id = request.GET.get("user_id")
    
    list_files = list(Files.objects.filter(
        context=context_id
    ).values(
        'id', 'document_id', 'document_type', 'user_id'
    ))
    
    return JsonResponse(list(list_files), safe=False)

@api_view(["GET", "POST"])
def create_admin_workspaces_contexts_files(request):
    user_id = request.GET.get("user_id")

    file = request.FILES.get('file', None)
    if file:
        filename = file.name
        file_extension = filename.split('.')[-1].lower()
        
        if file_extension not in ['']:
            return JsonResponse(
                {"error": "El archivo debe ser "},
                status=400
            )
    
    """
        Guarado a geonode
    """
            
    return JsonResponse( {"status": "ok"}, safe=False)


class ConsultaDocumentoView(APIView):

    def post(self, request):
        pregunta = request.data.get("pregunta")
        if not pregunta:
            return Response({"error": "Se requiere una pregunta."}, status=400)

        # Configurar LLM y embeddings
        llm = Ollama(model='llama3')
        embed_model = HuggingFaceEmbedding(model_name='sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2')

        # Conexión a vector store
        vector_store = PGVectorStore.from_params(
            database='tu_db',
            host='localhost',
            password='tu_password',
            port=5432,
            user='tu_usuario',
            table_name='documentos_vectorizados'
        )

        # Recuperar índice
        storage_context = StorageContext.from_defaults(vector_store=vector_store)
        index = VectorStoreIndex.from_vector_store(vector_store, storage_context=storage_context, embed_model=embed_model, llm=llm)

        # Ejecutar consulta
        query_engine = index.as_query_engine()
        respuesta = query_engine.query(pregunta)

        return Response({"respuesta": str(respuesta)}, status=200)


def cargar_documentos_desde_directorio(directorio: str) -> list[Document]:
    documentos = []

    for archivo in os.listdir(directorio):
        ruta = os.path.join(directorio, archivo)
        extension = archivo.lower().split('.')[-1]

        try:
            if extension == 'pdf':
                texto = extract_pdf_text(ruta)

            elif extension == 'txt':
                with open(ruta, 'r', encoding='utf-8') as f:
                    texto = f.read()

            elif extension == 'csv':
                df = pd.read_csv(ruta)
                texto = df.to_string(index=False)

            elif extension in ['xls', 'xlsx']:
                df = pd.read_excel(ruta)
                texto = df.to_string(index=False)

            elif extension == 'docx':
                doc = DocxDocument(ruta)
                texto = "\n".join([p.text for p in doc.paragraphs])

            else:
                print(f"Extensión no soportada: {archivo}")
                continue

            documentos.append(Document(text=texto, metadata={"nombre_archivo": archivo}))

        except Exception as e:
            print(f"Error procesando {archivo}: {e}")

    return documentos
